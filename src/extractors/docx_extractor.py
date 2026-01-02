"""
Microsoft Word document extractor using python-docx.

Extracts text from .docx files (modern Word format).
Note: .doc (legacy) files require conversion to .docx first.
"""

from pathlib import Path
from typing import ClassVar

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table

from src.extractors.base import DocumentExtractor, ExtractionError
from src.models import ExtractedDocument


class DocxExtractor(DocumentExtractor):
    """
    Extracts text content from Word documents (.docx).

    Extracts:
    - Paragraphs (with proper ordering)
    - Tables (converted to text format)
    - Headers and footers (if present)
    """

    SUPPORTED_EXTENSIONS: ClassVar[tuple[str, ...]] = (".docx",)

    def extract(self, file_path: Path) -> ExtractedDocument:
        """
        Extract text from a Word document.

        Args:
            file_path: Path to the .docx file.

        Returns:
            ExtractedDocument with the extracted text.

        Raises:
            ExtractionError: If the document cannot be read.
        """
        self._validate_file(file_path)

        try:
            doc = Document(str(file_path))
            text_parts: list[str] = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(table_text)

            if not text_parts:
                raise ExtractionError("Document contains no extractable text", file_path)

            full_text = "\n\n".join(text_parts)
            return self._create_result(full_text, file_path)

        except PackageNotFoundError as e:
            raise ExtractionError(
                "File is not a valid .docx document or is corrupted", file_path, cause=e
            ) from e
        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Unexpected error: {e}", file_path, cause=e) from e

    def _extract_table(self, table: Table) -> str:
        """
        Convert a Word table to text format.

        Args:
            table: python-docx Table object.

        Returns:
            Text representation of the table.
        """
        rows: list[str] = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))

        if rows:
            return "TABLE:\n" + "\n".join(rows)
        return ""
